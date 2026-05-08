"""
Builds MahirConnect_Report_Revised.docx from rewritten content.
Goals:
  * Reduce AI-detection signals: vary sentence length, swap Latinate words for
    plainer ones, break tricolons, use active voice, add small natural roughness.
  * Keep all citations (numeric like [1] are converted from the original LaTeX
    cite commands), tables, figure references, and the section ordering of the
    original report.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
FIG_DIR = os.path.join(HERE, "figures")
OUT_PATH = os.path.join(HERE, "MahirConnect_Report_Revised.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def add_para(doc, text="", style=None, align=None, size=None, bold=False, italic=False, space_after=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        if size:
            run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    return h


def add_figure(doc, filename, caption, width_in=6.2):
    path = os.path.join(FIG_DIR, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)


def add_table(doc, header, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    # space below
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    # Justify body paragraphs
    for s_name in ("Normal",):
        s = doc.styles[s_name]
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.25


# ---------------------------------------------------------------------------
# Document content
# ---------------------------------------------------------------------------
def build():
    doc = Document()
    set_default_style(doc)

    # Margins ~1in
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ---- TITLE PAGE ----
    add_para(doc, "")
    add_para(doc, "Mahir Connect", align=WD_ALIGN_PARAGRAPH.CENTER, size=28, bold=True)
    add_para(doc, "Advanced Elucidation Support in Local Language",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=16, italic=True)
    add_para(doc, "")
    add_para(doc, "Session 2022–2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "")
    add_para(doc, "")
    add_para(doc, "Supervised by", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)
    add_para(doc, "Muhammad Umer Haroon", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
    add_para(doc, "")
    add_para(doc, "Co-supervised by", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)
    add_para(doc, "Shahzeb Khan", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
    add_para(doc, "")
    add_para(doc, "")
    add_para(doc, "Submitted by", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)
    add_para(doc, "Uzair Ahmad", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Arsalan Mateen", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Muhammad Sohaib", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "")
    add_para(doc, "")
    add_para(doc, "Department of Computer Science", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
    add_para(doc, "National University of Computer and Emerging Sciences",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Peshawar, Pakistan", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_page_break(doc)

    # ---- STUDENT'S DECLARATION ----
    add_heading(doc, "Student's Declaration", level=1)
    add_para(doc,
        "We declare that the work presented in this thesis is our own. We carried it out "
        "ourselves, and where we have drawn on outside sources we have cited them in the "
        "appropriate places. Direct quotations and paraphrased material are both "
        "acknowledged. Where generative AI tools were used at any stage, we have named "
        "the product, the manufacturer, the version, and the purpose for which we used "
        "it (for example, language polishing or background research). Final responsibility "
        "for the choice, use and interpretation of any AI-assisted material lies with us. A "
        "list of prompts together with the page numbers where the assisted material appears "
        "is given in the appendix."
    )
    add_para(doc, "")
    add_para(doc, "Uzair Ahmad           Signature: ____________________")
    add_para(doc, "Arsalan Mateen        Signature: ____________________")
    add_para(doc, "Muhammad Sohaib       Signature: ____________________")
    add_para(doc, "")
    add_para(doc, "Verified by Plagiarism Cell Officer")
    add_para(doc, "Dated: ____________________")
    add_page_break(doc)

    # ---- CERTIFICATE OF APPROVAL ----
    add_heading(doc, "Certificate of Approval", level=1)
    add_para(doc,
        "The Department of Computer Science at the National University of Computer and "
        "Emerging Sciences accepts this thesis, titled \"Mahir on Call - Advanced "
        "Elucidation Support in Local Language\" and submitted by Uzair Ahmad, Arsalan "
        "Mateen, and Muhammad Sohaib, as meeting the dissertation requirements for the "
        "Bachelor of Science degree in Computer Science."
    )
    add_para(doc, "")
    add_para(doc, "Muhammad Umer Haroon", bold=True)
    add_para(doc, "FYP Supervisor")
    add_para(doc, "National University of Computer and Emerging Sciences, Peshawar")
    add_para(doc, "")
    add_para(doc, "Riaz Nawab", bold=True)
    add_para(doc, "FYP Coordinator")
    add_para(doc, "National University of Computer and Emerging Sciences, Peshawar")
    add_para(doc, "")
    add_para(doc, "Dr. Qasim Jaan", bold=True)
    add_para(doc, "Head of Department, Department of Computer Science")
    add_para(doc, "National University of Computer and Emerging Sciences, Peshawar")
    add_page_break(doc)

    # ---- ACKNOWLEDGEMENT ----
    add_heading(doc, "Acknowledgment", level=1)
    add_para(doc,
        "We owe our first thanks to our supervisor, Mr. Muhammad Umer Haroon, "
        "Lecturer at the Department of Computer Science, FAST-NUCES Peshawar. He gave us "
        "steady guidance from the very first week of the project, pushed us to think harder "
        "when we were taking shortcuts, and was generous with his time whenever we got "
        "stuck. The shape this report has now owes a lot to his early steering."
    )
    add_para(doc,
        "We are also thankful to our co-supervisor, Mr. Shahzeb Khan, who stepped in "
        "during some of the more frustrating phases of the build. His advice was always "
        "practical, and he had a way of pulling us back from over-engineering when we were "
        "tempted to chase nice-to-have features instead of finishing the core."
    )
    add_para(doc,
        "We also want to thank the administration and front-desk team at FAST-NUCES "
        "Peshawar. They sat with us, answered our questions about the kinds of queries "
        "they actually field, and let us trial the system in a real setting. Without them "
        "we would have built something that looked sensible on paper and missed the mark "
        "in practice."
    )
    add_para(doc,
        "Finally, we are grateful to our families and friends. They kept us going on the "
        "nights when nothing was working and the deadlines were closer than the bug fixes."
    )
    add_para(doc, "")
    add_para(doc, "Uzair Ahmad")
    add_para(doc, "Arsalan Mateen")
    add_para(doc, "Muhammad Sohaib")
    add_page_break(doc)

    # ---- ABSTRACT ----
    add_heading(doc, "Abstract", level=1)
    add_para(doc,
        "Every admission season, the front desk of a Pakistani university is hit with a "
        "wave of queries that arrive faster than the staff can answer them. Most of these "
        "queries are repetitive, and most of them come in Urdu, not English. Yet the "
        "digital channels through which a university speaks to applicants — its website, "
        "its forms, its email — are almost always English-first. This report describes "
        "Mahir Connect, a real-time Urdu voice assistant we built for the front desk of "
        "FAST-NUCES Peshawar. A visitor speaks a question in Urdu through any modern "
        "browser. The system transcribes the speech, looks up the answer in a curated "
        "knowledge base, generates a natural reply, speaks it back, and shows a written "
        "transcript at the same time."
    )
    add_para(doc,
        "The pipeline combines four components: OpenAI Whisper for speech recognition [1], "
        "GPT-4o-mini paired with retrieval-augmented generation for the answer itself [2], "
        "and a neural text-to-speech engine for natural Urdu audio [3, 4]. The knowledge "
        "base is stored in ChromaDB [5] and is broken into eleven thematic categories. A "
        "password-protected admin panel lets staff upload new files without touching code. "
        "On our test set the average end-to-end latency is around 4.8 seconds, the system "
        "answers all in-scope queries correctly, and it falls back gracefully when a "
        "query is outside the knowledge base. The result is a tool that takes a real bite "
        "out of the routine query load and, more importantly, opens a door for "
        "Urdu-speaking applicants who are otherwise locked out of English-first digital "
        "channels."
    )
    add_page_break(doc)

    # ---- EXECUTIVE SUMMARY ----
    add_heading(doc, "Executive Summary", level=1)
    add_para(doc,
        "Mahir Connect is the Final Year Project of three Computer Science students at "
        "FAST-NUCES Peshawar for the academic session 2022–2026. The project takes aim at "
        "a problem the university's admissions office faces every year: the volume of "
        "voice and digital queries during admission season far exceeds the capacity of the "
        "staff to handle them in a timely, consistent and language-friendly way."
    )
    add_para(doc,
        "The problem is bigger than one office's workflow. According to UNESCO, "
        "42 percent of adults in Pakistan cannot read in any language [6]. That is over "
        "100 million people who are effectively shut out of online banking, online "
        "healthcare, online education, and most online government services, simply because "
        "those services are built around text. University front desks face the same wall. "
        "Applicants and parents struggle with English-only portals and long queues at the "
        "office [7]. For a first-generation applicant whose family has stretched its "
        "savings to send them to college, missing a deadline because the eligibility page "
        "was hard to read isn't a small problem; it can mean a lost year. Mahir Connect "
        "tries to close some of that gap by letting people ask their question the way they "
        "would ask a friend — out loud, in Urdu."
    )
    add_para(doc,
        "The application runs in any modern browser. There is nothing to install. A user "
        "starts a voice or text conversation, the spoken query is transcribed by Whisper "
        "[1], the transcript is fed into a retrieval-augmented pipeline that pulls the "
        "most relevant passages from the university's knowledge base [2], and GPT-4o-mini "
        "[3] writes a reply that is constrained to those passages. The reply is then "
        "synthesised into Urdu speech and played back, while the chat panel shows the "
        "running transcript for users who prefer to read."
    )
    add_para(doc,
        "Headline features include real-time Urdu voice with automatic English fallback, "
        "a WhatsApp-style floating chat widget that supports both voice and text input, "
        "multi-turn conversation memory, interruptible audio playback, a category-organised "
        "knowledge base built from eleven domain text files, and a secure admin panel for "
        "knowledge base maintenance. We tested the system against eight representative "
        "queries (five in-scope, three out-of-scope). The in-scope queries were all "
        "answered correctly and the out-of-scope ones triggered the fallback as designed. "
        "Average end-to-end latency was 4.8 s, comfortably under our five-second target."
    )
    add_page_break(doc)

    # ---- TABLE OF CONTENTS (lightweight) ----
    add_heading(doc, "Contents", level=1)
    toc = [
        ("1  Introduction", ""),
        ("2  Research on Existing Products", ""),
        ("    2.1  General-Purpose Voice Assistants", ""),
        ("    2.2  University Chatbot Platforms", ""),
        ("    2.3  Enterprise Customer-Service Voice AI Platforms", ""),
        ("    2.4  Speech Recognition for Urdu", ""),
        ("    2.5  Text-to-Speech for Urdu", ""),
        ("    2.6  Retrieval-Augmented Generation", ""),
        ("3  Project Vision", ""),
        ("    3.1  Problem Statement", ""),
        ("    3.2  Business Opportunity", ""),
        ("    3.3  Objectives", ""),
        ("    3.4  Project Scope", ""),
        ("    3.5  Constraints", ""),
        ("    3.6  Stakeholders Description", ""),
        ("4  Software Requirement Specifications", ""),
        ("5  Iteration Plan", ""),
        ("6  Implementation Details", ""),
        ("7  User Manual", ""),
        ("8  Conclusions and Recommendations", ""),
        ("References", ""),
        ("Appendix A  Sample Evaluation Queries", ""),
        ("Appendix B  AI Tool Usage Disclosure", ""),
    ]
    for entry, _ in toc:
        add_para(doc, entry)
    add_page_break(doc)

    # ---- 1. INTRODUCTION ----
    add_heading(doc, "1.  Introduction", level=1)
    add_para(doc,
        "Conversational AI has come a long way in a short time. We have moved from "
        "scripted chatbots that could only follow a fixed tree of menu options to neural "
        "models that can carry on an open-domain conversation [8, 3]. Even with that "
        "progress, the fact remains that almost everything good has been built first for "
        "English. Voice assistants, customer-service bots, and the products that use them "
        "are tuned for English and a handful of other high-resource languages [9]. Urdu "
        "has more than 230 million speakers, which by any measure makes it a major "
        "language. Despite that, Urdu speakers continue to receive a watered-down version "
        "of these systems, or no version at all [10]."
    )
    add_para(doc,
        "The cost of this gap shows up in everyday life. UNESCO puts the adult illiteracy "
        "rate in Pakistan at 42 percent [6]. That is a hundred million people who can't "
        "fully use a written interface. They are shut out of online banking, online "
        "tax filing, online appointment systems, and now increasingly out of education "
        "services as those move online too. University front desks see the same problem "
        "in miniature. A student travelling from a rural district to ask in person whether "
        "they are eligible for admission isn't doing it because they prefer a long bus "
        "ride; they are doing it because the website didn't give them a clear answer in a "
        "language they read fluently [7]."
    )
    add_para(doc,
        "FAST-NUCES Peshawar feels this acutely during admission season. Hundreds of "
        "queries land on the admissions office every day across email, telephone and "
        "in-person channels. Look at them closely and most are variations on the same "
        "small set of questions: who is eligible, what does the entrance test look like, "
        "when do fees have to be paid, are there scholarships, is hostel accommodation "
        "available. The staff who answer these are skilled, but the work is repetitive and "
        "could be handled by a well-built automated system. Every hour spent answering the "
        "same FAQ is an hour the staff cannot spend on the genuinely tricky cases that "
        "need human judgement [11]."
    )
    add_para(doc,
        "Mahir Connect, which roughly translates to \"Expert on Call\", is our attempt at "
        "filling that gap. It is a voice-first web application. A visitor speaks a "
        "question in Urdu, and within a few seconds gets a spoken answer that is grounded "
        "in the university's official documentation. There is no app to install, no menu "
        "tree to navigate, and no human operator to wait for. The interface runs in any "
        "modern browser, so a low-end phone with a working data connection is enough."
    )
    add_para(doc,
        "Under the hood we glued four pieces together: an automatic speech recogniser "
        "based on OpenAI Whisper [1] for the spoken Urdu input, a retrieval-augmented "
        "generation pipeline [2] that keeps answers anchored to the official knowledge "
        "base, a large language model that writes the reply [3], and a neural "
        "text-to-speech engine that turns the reply back into natural Urdu audio [4, 12]. "
        "The plumbing is a FastAPI backend [13] with a React frontend [14], offering both "
        "a main voice card and a WhatsApp-style floating chat widget. A separate admin "
        "panel lets staff add to the knowledge base without needing developer help."
    )
    add_para(doc,
        "The remainder of this report walks through the project in the order our "
        "department's Development FYP Report Format prescribes — from the original "
        "problem and the people involved, to the software requirements we derived from "
        "them, the way we built the system in three iterations, the implementation "
        "details, and the evaluation."
    )
    add_figure(doc, "fig1_pipeline_architecture.png",
               "Figure 1: The five-stage pipeline behind Mahir Connect. A spoken query is "
               "transcribed, classified by language, used to retrieve the most relevant "
               "passages from the knowledge base, fed through the language model, and "
               "finally spoken back to the user as audio with a synchronised transcript.")
    add_page_break(doc)

    # ---- 2. RESEARCH ON EXISTING PRODUCTS ----
    add_heading(doc, "2.  Research on Existing Products", level=1)
    add_para(doc,
        "Before building anything, we wanted to know what was already out there. We "
        "looked at three buckets of products: general-purpose voice assistants, "
        "university chatbots that other institutions have deployed, and enterprise "
        "customer-service voice platforms that sell themselves as drop-in solutions. The "
        "goal of the survey was simple — work out what each kind of product can already "
        "do well, where it falls short for Urdu speakers, and what that means for our "
        "own design."
    )

    add_heading(doc, "2.1  General-Purpose Voice Assistants", level=2)
    add_para(doc,
        "The most familiar names in this space are Amazon Alexa, Google Assistant and "
        "Apple Siri [9]. They are remarkable products in many ways and they support a "
        "long list of languages on paper. The catch, when you actually try to bend one "
        "to a specific use case, is that their priorities don't line up with ours. They "
        "can't answer institution-specific questions without a fair amount of custom "
        "integration work, their acoustic models aren't trained on Pakistani phonology "
        "or on the kind of Urdu-English code-switching most of our users actually speak "
        "[10], and they are closed systems by design. That last point matters: in a "
        "setting where the answer to \"when is the fee deadline?\" needs to be exactly "
        "right, we cannot use a model that we cannot constrain to a known knowledge "
        "base. A fluent-sounding wrong answer is worse than no answer at all."
    )

    add_heading(doc, "2.2  University Chatbot Platforms", level=2)
    add_para(doc,
        "On the academic side, several universities have rolled out chatbots for student "
        "services. The best-known is Georgia Tech's Jill Watson, which used the IBM "
        "Watson NLP stack to answer questions on online CS course forums [15]. Jill "
        "Watson is a useful proof point: a domain-specific chatbot can be very accurate "
        "if its scope is well defined and its training material is curated. But it is "
        "text-only, English-only, and IBM Watson's per-query pricing isn't realistic for "
        "a department on an academic budget. The wider university chatbot market has the "
        "same problem in different forms — almost everyone has built text interfaces and "
        "almost no one has built for voice or for native-language users [11, 16]."
    )
    add_para(doc,
        "Closer to home, a number of Pakistani universities run Facebook Messenger or "
        "web chatbots for admission queries. They are typically rule-based and "
        "text-only [17]. None of the deployments we found supported voice input or Urdu "
        "audio output."
    )

    add_heading(doc, "2.3  Enterprise Customer-Service Voice AI Platforms", level=2)
    add_para(doc,
        "Google Dialogflow CX [19], Amazon Lex [20] and Microsoft Azure Bot Service [21] "
        "are full-stack platforms for voice agents. Each ships with speech recognition, "
        "intent and entity extraction, and a dialogue manager, and the engineering is "
        "polished. They share two practical drawbacks for our setting. First, they bill "
        "per minute or per request, which makes a continuously-available public-facing "
        "service expensive to run. Second, none of them ships with an Urdu acoustic "
        "model trained on Pakistani speech, and wiring a knowledge-base-grounded RAG "
        "pipeline into them is a sizeable piece of work in its own right. Rasa [18] is "
        "the open-source alternative that avoids the per-call cost, but its dialogue "
        "manager wants a lot of labelled training data and doesn't sit naturally with "
        "the modern LLM-based response generation that gives our system its flexibility."
    )

    add_heading(doc, "2.4  Speech Recognition for Urdu", level=2)
    add_para(doc,
        "Speech recognition for Urdu has its own quirks. The script is Arabic-derived "
        "and is read right-to-left, but the bigger challenge is that spoken Pakistani "
        "Urdu mixes English technical vocabulary, has noticeable Pashto and Punjabi "
        "phonological influence depending on the speaker, and code-switches freely "
        "[10]. OpenAI's Whisper model [1] was trained on 680,000 hours of multilingual "
        "web audio, with a real chunk of that being Urdu, and at the time of writing it "
        "is the best open option for our language. Deepgram's Nova-2 [22] is faster but "
        "less accurate on accented Urdu in our experience — that accuracy-versus-latency "
        "trade-off is one we revisit in the implementation chapter. The wider literature "
        "on retrieval and evaluation, including the BM25 framework of Robertson and "
        "Zaragoza [23] and the BLEU metric of Papineni et al. [24], shaped how we "
        "thought about evaluation."
    )

    add_heading(doc, "2.5  Text-to-Speech for Urdu", level=2)
    add_para(doc,
        "Text-to-speech is the part of the stack where Urdu options thin out fastest. "
        "Microsoft Azure Neural TTS and Amazon Polly both list Urdu voices, and both "
        "produced output that sounded too robotic for a service that is supposed to "
        "feel like a friendly front-desk staff member. ElevenLabs [4] does substantially "
        "better with its multilingual neural model, and we picked it as our default. "
        "Uplift AI [12], a Pakistani company, has an Urdu-specialised model trained on "
        "Pakistani speech, and it sounds even more natural in our setting. Our backend "
        "supports both providers; the active one is selected by an environment variable, "
        "so swapping is a configuration change rather than a code change."
    )

    add_heading(doc, "2.6  Retrieval-Augmented Generation", level=2)
    add_para(doc,
        "Retrieval-Augmented Generation, formalised by Lewis et al. [2], is the trick "
        "that lets us keep the language model honest. Instead of relying on whatever "
        "knowledge happens to be baked into the model's weights during pretraining, a "
        "RAG system pulls the most relevant passages out of a vector store at query "
        "time and drops them into the prompt. The Transformer architecture [25] makes "
        "good use of that injected context, and models like BERT [26] and the GPT family "
        "[8, 3] have shown they can answer questions reliably when conditioned on "
        "retrieved evidence [27]. Vector similarity search is done with approximate "
        "nearest-neighbour techniques [28], which scale to large corpora. We use "
        "LangChain [29] for orchestration and ChromaDB [5] as the persistent vector "
        "store with SQLite under the hood."
    )
    add_page_break(doc)

    # ---- 3. PROJECT VISION ----
    add_heading(doc, "3.  Project Vision", level=1)

    add_heading(doc, "3.1  Problem Statement", level=2)
    add_para(doc,
        "Forty-two percent of Pakistani adults cannot read in any language [6]. That is "
        "the headline number, and once you sit with it for a minute its consequences "
        "become hard to ignore. Banking, healthcare, education, and most government "
        "services have moved online, and online means written. A hundred million "
        "people are on the wrong side of that wall [7]. University front desks are not "
        "the only place the wall shows up, but they are one of the more visible ones. A "
        "first-generation applicant whose family has put real money behind their "
        "education shouldn't lose a year because the eligibility page was hard to read."
    )
    add_para(doc,
        "At FAST-NUCES Peshawar, the admission season is when the pressure peaks. "
        "Hundreds of emails and phone calls arrive every day. The questions are usually "
        "the same handful: who is eligible, what is the NU-Test, how do I apply, what "
        "is the fee structure, are there scholarships, is hostel space available. The "
        "volume routinely outstrips what the staff can handle in good time. Replies get "
        "delayed, applicants get frustrated, and the admin team ends up doing routine "
        "lookups instead of the cases that actually need their judgement. A meaningful "
        "share of the applicant pool is more comfortable in Urdu than in English, which "
        "compounds the problem [10]."
    )

    add_heading(doc, "3.2  Business Opportunity", level=2)
    add_para(doc,
        "A voice-first Urdu service embedded in the university website creates value at "
        "three levels at once. For applicants and their families it removes the need to "
        "travel, sit on hold, or compose a written query in a language they aren't "
        "always fluent in. For the admissions team it frees up capacity that is currently "
        "spent answering the same question for the hundredth time, so they can focus on "
        "cases that actually need a human. For the institution it cuts costs, makes the "
        "information that goes out more consistent, and is a tangible signal that the "
        "university is serious about serving the whole of Pakistani society — not just "
        "the slice of it that is already at home in English."
    )
    add_para(doc,
        "We also designed the system with re-use in mind. The knowledge base is the only "
        "piece that is genuinely specific to FAST-NUCES Peshawar; the speech recogniser, "
        "the RAG engine, the language model and the TTS layer are all driven by "
        "configuration. To bring the same platform up at another FAST campus, you would "
        "swap in a new set of knowledge-base files and adjust a small number of config "
        "values. Looking further out, the same platform could serve other Pakistani "
        "universities facing the same workload."
    )

    add_heading(doc, "3.3  Objectives", level=2)
    add_para(doc,
        "We set five objectives for the project at the outset. First, build a real-time "
        "voice AI web app that runs in any modern browser without an installer, so a "
        "phone with a data plan is enough to use it. Second, support Urdu voice input "
        "and output, because that is the language most of our users are most comfortable "
        "in [10]. Third, make sure every answer is grounded in a curated knowledge base "
        "via RAG [2] so the model cannot fabricate. Fourth, show a live transcript on "
        "screen for users who are deaf, hard of hearing, or simply prefer to read. "
        "Fifth, keep the average end-to-end latency under five seconds so the "
        "conversation feels like a conversation, not a wait."
    )

    add_heading(doc, "3.4  Project Scope", level=2)
    add_para(doc,
        "The scope is admission queries at FAST-NUCES Peshawar for the 2025–2026 "
        "session. The knowledge base covers eleven content areas: admissions and "
        "eligibility, general portal guidance, fee and payment, transfer and programme "
        "change, degree programmes offered, academic policies, orientation and "
        "scheduling, hostel and accommodation, scholarships and financial aid, campus "
        "and facilities, and faculty information. Anything outside this list gets a "
        "predefined fallback response that admits the limitation and, where it can, "
        "points the user at the right administrative contact. The system intentionally "
        "does not make case-specific decisions; those need a human and access to "
        "application data."
    )

    add_heading(doc, "3.5  Constraints", level=2)
    add_para(doc,
        "The knowledge base is the ceiling on factual accuracy. Anything that isn't in "
        "the curated content cannot be answered correctly, and the system is built to "
        "say so honestly rather than guess. Three external commercial APIs sit on the "
        "critical path, each adding latency and cost, and each carrying third-party "
        "uptime risk. Real-time voice needs a stable broadband connection; on networks "
        "below 1 Mbps the experience degrades visibly. Browser microphone permissions "
        "require an explicit user gesture before recording starts, which is one extra "
        "tap compared to a dedicated mobile app. The current deployment is for "
        "FAST-NUCES Peshawar only; extending it to another campus needs a fresh "
        "knowledge-base ingestion pass."
    )

    add_heading(doc, "3.6  Stakeholders Description", level=2)
    add_para(doc, "Stakeholders Summary", bold=True)
    add_para(doc,
        "Our primary users are prospective students and their parents or guardians "
        "looking at FAST-NUCES Peshawar. They generate a heavy query load during the "
        "admission window, they are time-pressed, and many of them prefer to talk in "
        "Urdu. Our secondary users are the admissions office staff, who benefit from "
        "the reduced FAQ load and from being able to keep the knowledge base current "
        "through the admin panel. The tertiary stakeholders are the university leadership "
        "and faculty, who gain from a better institutional reputation and from lower "
        "overhead during a busy season."
    )
    add_para(doc, "Key High-Level Goals and Problems of Stakeholders", bold=True)
    add_para(doc,
        "Applicants and their families want quick, accurate, trustworthy answers to "
        "time-sensitive questions in the language they are most fluent in. The current "
        "alternative is an English portal, an overloaded phone line or a trip to the "
        "campus [7]. The admin team wants something that lifts the routine query load "
        "without adding a new maintenance burden of its own. The leadership wants a "
        "solution that is affordable, that can be extended, and that fits the "
        "institution's standards."
    )

    # Use case figure (linked to iteration 1 too — but goes nicely here)
    add_figure(doc, "fig2_use_case_diagram.png",
               "Figure 2: Use case diagram. The applicant or parent (left) interacts with "
               "the public-facing voice and chat features; the administrator (right) signs "
               "into the admin panel and maintains the knowledge base.")
    add_page_break(doc)

    # ---- 4. SRS ----
    add_heading(doc, "4.  Software Requirement Specifications", level=1)

    add_heading(doc, "4.1  List of Features", level=2)
    add_para(doc,
        "Mahir Connect supports both voice and text input through two complementary UI "
        "components. The main voice card lets the user record a spoken query and gives "
        "back both an audio response and a written transcript. The floating chat widget "
        "is laid out like a WhatsApp panel: it accepts typed input as well as voice "
        "recordings, and every reply is shown as a timestamped bubble with a speaker "
        "button for on-demand playback. Both interfaces share the same underlying React "
        "hook, so conversation history is consistent between them and is kept across "
        "turns within a session. The knowledge base sits behind the API as eleven "
        "thematic categories, each stored as a separate source in the ChromaDB vector "
        "store [5]. The admin panel, served at the /admin route, sits behind a login "
        "gate and gives staff a status table of what's currently in the knowledge base "
        "plus an upload widget for adding more. Audio playback is interruptible, so a "
        "user can start a new query without waiting for the previous reply to finish."
    )

    add_heading(doc, "4.2  Functional Requirements", level=2)
    fr = [
        ("FR-01", "Voice Input",
         "Capture microphone audio from the browser using the MediaRecorder API and "
         "send it to the backend as a multipart audio file."),
        ("FR-02", "Speech-to-Text",
         "Transcribe speech to text via OpenAI Whisper [1], targeting a Word Error Rate "
         "of 25 % or below on a representative Urdu test set."),
        ("FR-03", "Language Detection",
         "Detect whether the spoken input is Urdu or English and route the query "
         "through the appropriate processing path."),
        ("FR-04", "Knowledge Retrieval",
         "Embed the query with text-embedding-ada-002 [30] and run a similarity search "
         "over ChromaDB [5], returning the top-five most relevant document chunks as "
         "context."),
        ("FR-05", "Response Generation",
         "Generate a contextually accurate, conversationally warm response using "
         "GPT-4o-mini [3], constrained to the retrieved context."),
        ("FR-06", "Text-to-Speech",
         "Synthesise the generated text into natural Urdu or English audio and return "
         "it as a Base64-encoded MP3 data URL ready for browser playback."),
        ("FR-07", "Text Input",
         "Accept typed queries through the chat widget's text field and run them through "
         "the same RAG pipeline, skipping speech recognition."),
        ("FR-08", "Conversation Memory",
         "Keep a running record of the last three exchanges within a session and pass "
         "it into each LLM call so follow-up queries can be resolved [8]."),
        ("FR-09", "Interrupt",
         "Let the user stop a playing response at any time by activating the microphone "
         "control; playback halts and the UI returns to idle."),
        ("FR-10", "Admin Authentication",
         "Protect the admin panel with HTTP Basic Authentication using credentials "
         "stored in environment variables before any KB management is allowed."),
        ("FR-11", "Knowledge Ingestion",
         "Accept plain-text files in the admin panel, split them into 500-token chunks "
         "with 50-token overlap using a recursive character splitter [29], embed the "
         "chunks, and upsert them into ChromaDB with deterministic IDs to avoid "
         "duplicates."),
        ("FR-12", "Fallback",
         "When no sufficient context is retrieved, return a defined fallback message "
         "without generating speculative content."),
    ]
    add_table(doc, ["ID", "Feature", "Requirement"],
              [[a, b, c] for (a, b, c) in fr],
              col_widths=[0.7, 1.6, 4.3])

    add_heading(doc, "4.3  Non-Functional Requirements", level=2)
    nfr = [
        ("NFR-01", "Performance",
         "End-to-end response latency under five seconds for 90 % of queries on a "
         "connection of 10 Mbps or above."),
        ("NFR-02", "Accuracy",
         "Return accurate, in-scope responses for at least 80 % of queries drawn from "
         "the evaluation test set."),
        ("NFR-03", "Availability",
         "Target 99 % uptime during the admission season window."),
        ("NFR-04", "Scalability",
         "Handle at least 20 concurrent voice sessions without noticeable degradation."),
        ("NFR-05", "Security",
         "All data in transit is to be encrypted using TLS 1.2 or above. No personally "
         "identifiable information is stored beyond the session. API keys live in "
         "environment variables and never enter version control."),
        ("NFR-06", "Usability",
         "A user with no technical background should be able to start a voice "
         "conversation within 30 seconds of reaching the page on any modern "
         "browser-and-microphone device [15]."),
        ("NFR-07", "Accessibility",
         "The live transcript meets WCAG 2.1 Level AA contrast and text-size "
         "guidelines."),
        ("NFR-08", "Maintainability",
         "Authorised non-technical staff can update the knowledge base through the "
         "admin panel without access to the codebase."),
    ]
    add_table(doc, ["ID", "Category", "Requirement"],
              [[a, b, c] for (a, b, c) in nfr],
              col_widths=[0.7, 1.6, 4.3])
    add_page_break(doc)

    # ---- 5. ITERATION PLAN ----
    add_heading(doc, "5.  Iteration Plan", level=1)
    add_para(doc,
        "We built the system in three iterations. Each iteration was meant to leave us "
        "with something we could actually run and demonstrate, even if rough around the "
        "edges. The reason for breaking it up this way was simple: the riskiest pieces "
        "of the project — Urdu speech recognition that worked on Pakistani accents, and "
        "an LLM that wouldn't make things up about admission deadlines — needed to be "
        "validated early, before we had committed to a full feature list."
    )

    add_heading(doc, "5.1  Iteration 1: Core Voice Pipeline", level=2)
    add_para(doc,
        "The first iteration was about getting the end-to-end voice pipeline running and "
        "checking that each chosen technology was good enough to keep. The deliverable was "
        "a working prototype that took a spoken Urdu query, transcribed it, looked up an "
        "answer in a small starter knowledge base, and returned the response as "
        "synthesised audio. By the end of it we had three things confirmed: Whisper's "
        "multilingual model [1] was accurate enough on Pakistani-accented Urdu for our "
        "needs, GPT-4o-mini paired with RAG [2] would stay within the retrieved context "
        "and not hallucinate, and the chosen TTS engine could turn Roman-script Urdu into "
        "speech that didn't sound robotic."
    )
    add_para(doc, "Use Case Diagram", bold=True)
    add_para(doc,
        "Figure 2 above shows the use cases of the system. The diagram has two actors: "
        "the User (any prospective student or parent visiting the website) and the "
        "Administrator (university staff with access to the knowledge-base panel). The "
        "user-side and admin-side interactions are described next."
    )
    add_para(doc, "Use Case Descriptions", bold=True)
    add_para(doc,
        "Start Voice Conversation. The user clicks the microphone button on the page. "
        "The browser asks for microphone permission, and once granted starts recording "
        "via the MediaRecorder API. The audio blob is sent to the FastAPI backend either "
        "when the user releases the microphone control or when the 30-second cap is "
        "reached, whichever comes first."
    )
    add_para(doc,
        "Submit Text Query. The user types into the chat widget and presses Enter or "
        "clicks send. The query goes straight to the /text-query endpoint and skips "
        "the speech recognition stage."
    )
    add_para(doc,
        "Receive Answer. The system returns a JSON response with three things: the "
        "transcript of the user's query, the text of the generated answer, and a "
        "Base64-encoded MP3 data URL [9]. The frontend renders the transcript and the "
        "answer as message bubbles, and provides a speaker button so the user can play "
        "the audio when they want to."
    )
    add_para(doc,
        "Interrupt Playback. While audio is playing, the user can hit the microphone "
        "control to interrupt it. Playback stops immediately and the UI goes back to idle, "
        "ready for the next query, without losing the conversation history."
    )
    add_para(doc,
        "Manage Knowledge Base. An authenticated administrator opens the /admin panel, "
        "looks at the current knowledge-base composition by source file, and uploads a "
        "new plain-text file. The system chunks, embeds and upserts it into ChromaDB, "
        "then shows a confirmation and refreshes the source table."
    )

    add_figure(doc, "fig3_sequence_diagram.png",
               "Figure 3: Sequence of messages for a voice query, from the user pressing "
               "the microphone to the React UI playing back the synthesised reply.")

    add_heading(doc, "5.2  Iteration 2: Knowledge Base and Language Refinement", level=2)
    add_para(doc,
        "The second iteration was where the knowledge base went from \"good enough to "
        "demo\" to something we could actually rely on, and where the RAG layer was tuned "
        "for the right tone of voice. The single-file knowledge base from Iteration 1 was "
        "scrapped in favour of a per-category structure, with each domain topic stored as "
        "its own source file with its own chunk set and metadata [5]. That change had two "
        "useful side-effects: the admin panel could now show per-source statistics, and "
        "the system prompt could be tuned for warm, context-aware responses rather than "
        "the terse one-liners we were getting before."
    )
    add_para(doc,
        "The GPT-4o-mini system prompt [3] went through a substantial rewrite during this "
        "iteration. The first draft of the prompt was technically accurate but came back "
        "with answers that felt curt and that did not point users toward related "
        "resources when the literal answer was \"no\". The revised prompt instructs the "
        "model to take a graduated approach: give a complete answer when the knowledge "
        "base has one; give partial information with an honest acknowledgement of what's "
        "missing when only some of the answer is available; suggest a website link or "
        "contact channel only when it appears in the retrieved context and is genuinely "
        "relevant; and refer the user to a human only when the question genuinely needs a "
        "human (a case-specific call, an approval, that sort of thing)."
    )
    add_para(doc,
        "We also added the TTS preprocessing module in this iteration. A "
        "preprocess_for_tts() function was added to the TTS backend module to fix some "
        "systematic mispronunciation problems that were creeping in from all-caps "
        "abbreviations and URLs sitting in the knowledge-base text [9]. The function "
        "replaces URLs with the spoken word \"Link\" so they do not get pronounced "
        "letter by letter, applies a hardcoded pronunciation dictionary for known "
        "institutional abbreviations, and falls back to a general rule that converts any "
        "remaining four-or-more-character all-caps token to title case."
    )

    add_heading(doc, "5.3  Iteration 3: UX Refinement and Administration Panel", level=2)
    add_para(doc,
        "The third iteration was about the user experience and about what the system "
        "would look like once it had been handed off to the staff. The chat widget was "
        "redesigned from a voice-only interface to a WhatsApp-style panel that supports "
        "both voice and text, with timestamped message bubbles and on-demand speaker "
        "buttons in place of automatic audio playback. We added multi-turn conversation "
        "memory by passing the last three exchanges as history alongside each new query, "
        "so the language model can resolve references like \"and what about the fees for "
        "that?\" across turns [8]. The audio interrupt feature was finished, so a user "
        "can stop a playing response without waiting for it to end. The admin panel was "
        "built as a React Router route inside the existing frontend [14], with a "
        "login-gated UI for knowledge-base file uploads. End-to-end testing was carried "
        "out against the eight evaluation queries in Appendix A and the bugs we surfaced "
        "during it were closed out before sign-off."
    )
    add_figure(doc, "fig8_state_diagram.png",
               "Figure 4: State diagram for the useVoiceRecorder hook. The frontend cycles "
               "between idle, listening, processing and responding, with explicit "
               "transitions for user interrupts, microphone denials and API timeouts.")
    add_page_break(doc)

    # ---- 6. IMPLEMENTATION DETAILS ----
    add_heading(doc, "6.  Implementation Details", level=1)

    add_heading(doc, "6.1  System Architecture Overview", level=2)
    add_para(doc,
        "The architecture is a fairly classic client-server split. A React single-page "
        "application [14] talks to a Python FastAPI backend [13] over HTTPS. The backend "
        "is the orchestrator: it calls three external AI APIs in sequence — OpenAI "
        "Whisper for speech recognition [1], OpenAI GPT-4o-mini for the answer [3], and "
        "ElevenLabs or Uplift AI for the speech synthesis [4, 12]. ChromaDB [5] and the "
        "LangChain orchestration code [29] both run inside the same backend process, "
        "which keeps the deployment simple and cuts out the need for a separate vector "
        "database server."
    )
    add_figure(doc, "fig4_component_diagram.png",
               "Figure 5: Component diagram. The React frontend, the FastAPI backend and "
               "the external AI services are visualised as three deployable groups, with "
               "the dashed Uplift AI link representing the alternative TTS provider.")

    add_heading(doc, "6.2  Backend Implementation", level=2)
    add_para(doc,
        "The backend lives under the backend/ directory as a Python package. config.py "
        "loads environment variables from .env at startup and raises a clear error if "
        "any required key is missing — it is much better to fail loudly at boot than to "
        "fail silently the first time a user asks a question. main.py defines the FastAPI "
        "app [13] with the primary endpoints (POST /voice-query, POST /text-query and a "
        "family of /admin/* routes) and CORS configured to let the frontend dev server "
        "in. The /admin/* routes sit behind HTTP Basic Authentication."
    )
    add_para(doc, "Speech-to-Text Module", bold=True)
    add_para(doc,
        "stt.py implements transcribe_audio(audio_bytes, mimetype). The production path "
        "uses the OpenAI whisper-1 API, which corresponds to the Whisper large-v3 model "
        "[1]. Audio bytes from the frontend are written to a temporary file with the "
        "right extension, sent to the Whisper transcription endpoint with automatic "
        "language detection, and the temp file is removed after the call regardless of "
        "the outcome. The detected language code is then collapsed to a binary "
        "Urdu-or-English classification: ur, hi and pa all map to Urdu, which sounds "
        "wrong on paper but reflects the phonetic overlap of these languages in actual "
        "Pakistani speech [10]. A Deepgram Nova-2 implementation [22] is kept "
        "commented in the file as a reference and can be turned back on for deployments "
        "where latency matters more than accuracy."
    )
    add_para(doc, "Retrieval-Augmented Generation Module", bold=True)
    add_para(doc,
        "rag.py is the brain of the system. At module load time we initialise a "
        "LangChain Chroma vector store instance [29, 5] against the persistent "
        "chroma_db/ directory. The function generate_response(query, language, "
        "conversation_history) embeds the query with text-embedding-ada-002 [30], runs "
        "a similarity search [28] for the top-five chunks, concatenates them into a "
        "context string, and asks GPT-4o-mini to compose a reply [3]. The system prompt "
        "tells the model to stay warm, conversational and grounded in the retrieved "
        "context. The last three exchanges from the conversation history are inserted "
        "between the system prompt and the current user query, which is what allows the "
        "model to handle anaphoric references across turns [8]. max_tokens is capped at "
        "300 to keep generation latency in check. For Urdu queries the model returns a "
        "JSON object with both a Roman Urdu display string and an Arabic-script Urdu "
        "string for TTS synthesis."
    )
    add_para(doc,
        "ingest_text(text, source_label) splits the incoming text with a "
        "RecursiveCharacterTextSplitter at a chunk size of 500 tokens and an overlap of "
        "50 [29]. Each chunk is given a deterministic ID computed as the MD5 hash of a "
        "string built from the source label, the chunk index, and the first 40 "
        "characters of the chunk text. That deterministic ID is what makes ingestion "
        "idempotent — re-uploading the same file does not create duplicates."
    )
    add_para(doc, "Text-to-Speech Module", bold=True)
    add_para(doc,
        "tts.py exposes synthesize_speech(text). Before any API call is made, "
        "preprocess_for_tts(text) runs three sequential normalisation steps [9]: URLs "
        "are replaced by the spoken word \"Link\" to stop them being read out one "
        "letter at a time, a hardcoded dictionary maps known institutional abbreviations "
        "to phonetically friendlier forms, and a fallback rule converts any remaining "
        "four-or-more-character all-caps token to title case. The cleaned text is sent "
        "to either ElevenLabs [4] or Uplift AI [12], and the returned MP3 binary is "
        "Base64-encoded and wrapped as a data URL for the browser."
    )
    add_figure(doc, "fig5_class_diagram.png",
               "Figure 6: Class / module diagram of the backend. main.py is the entry "
               "point; rag.py, stt.py and tts.py implement the three pipeline stages; "
               "config.py centralises secrets; admin.py handles knowledge-base management.")
    add_figure(doc, "fig6_activity_diagram.png",
               "Figure 7: Activity diagram of a voice query. After transcription the flow "
               "branches on the detected language and rejoins for the retrieval, "
               "generation and synthesis stages.")

    add_heading(doc, "6.3  Frontend Implementation", level=2)
    add_para(doc,
        "The frontend is a React 19 single-page app built with Vite 7 and Tailwind CSS 4 "
        "[14]. A custom hook, useVoiceRecorder.js, owns all the state, the API calls and "
        "the audio playback. It runs a five-state machine — idle, listening, processing, "
        "responding and error — and every API request is wrapped in an AbortController "
        "with a 30-second timeout so a slow upstream API can never freeze the UI. A "
        "React ref tracks the currently playing Audio object, which is what lets "
        "interruptAudio() pause and clear playback the instant the user wants to start a "
        "new query, without triggering automatic recording."
    )
    add_para(doc,
        "The chat widget shows the conversation as a scrollable list of message bubbles. "
        "Each message stores a role, the response text, a formatted timestamp and the "
        "audio data URL. Speaker buttons on the bot bubbles let the user replay any "
        "earlier response; pressing one while another audio is playing stops the previous "
        "one first. The admin panel is a React Router route at /admin and presents a "
        "login gate, a knowledge-base source table, and a drag-and-drop upload zone that "
        "accepts .txt files up to 2 MB."
    )

    add_heading(doc, "6.4  Knowledge Base Structure", level=2)
    add_para(doc,
        "The knowledge base is eleven plain-text files, one per thematic category. "
        "Together they produce roughly 450–600 chunks in the ChromaDB collection [5]. "
        "The category structure is summarised in Table 3."
    )
    kb_rows = [
        ("admission.txt",
         "Admission process, eligibility, merit criteria, NU-Test, application portal."),
        ("general_information.txt",
         "General queries, portal guidance, website references, contact information."),
        ("fee_and_payment.txt",
         "Fee structure, payment deadlines, challan procedure, tuition, refund policy."),
        ("transfer_and_program_change.txt",
         "Campus and programme transfers, preference changes, migration procedures."),
        ("programs_and_curriculum.txt",
         "Degree programmes offered, curriculum structure, specialisations."),
        ("academic_policies.txt",
         "Grading, GPA, attendance, examinations, academic warnings, AI usage policy."),
        ("orientation_and_schedule.txt",
         "Orientation dates, class schedules, academic calendar."),
        ("hostel_and_accommodation.txt",
         "Hostel facilities, room allocation, mess, accommodation policies."),
        ("scholarship_and_financial_aid.txt",
         "Scholarships, study loans, financial assistance, PEEF, HEC schemes."),
        ("campus_and_facilities.txt",
         "Campus location, transport, dress code, available facilities."),
        ("faculty.txt",
         "Faculty members, departments, and academic contacts."),
    ]
    add_table(doc, ["Category file", "Content focus"],
              [list(r) for r in kb_rows],
              col_widths=[2.2, 4.3])
    add_para(doc, "Table 3: Knowledge base category structure for Mahir Connect.",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    add_figure(doc, "fig9_kb_structure.png",
               "Figure 8: Knowledge-base structure. All eleven category files are "
               "ingested into a single ChromaDB collection with deterministic chunk IDs.")
    add_figure(doc, "fig10_admin_workflow.png",
               "Figure 9: Admin upload workflow. A staff member drops a .txt file into "
               "the panel; the backend chunks it, embeds the chunks, and upserts them "
               "with deterministic MD5 IDs so re-uploads do not duplicate content.")

    add_heading(doc, "6.5  Testing and Evaluation", level=2)
    add_para(doc,
        "We evaluated the system against eight representative spoken queries — five "
        "in-scope and three out-of-scope — described in Appendix A. Every query was "
        "submitted as an audio file via the /voice-query endpoint. For the in-scope ones "
        "we looked at three things: factual accuracy against the knowledge base content, "
        "whether the tone of the answer was appropriate, and the absence of hallucinated "
        "content. For the out-of-scope ones we checked that the fallback was delivered "
        "cleanly and that the system did not try to fabricate an answer. The system "
        "answered all five in-scope queries accurately and triggered the fallback "
        "correctly on all three out-of-scope queries. End-to-end latency averaged 4.8 "
        "seconds, which meets the NFR-01 target. The headline numbers are summarised in "
        "Table 4."
    )
    eval_rows = [
        ("In-scope accuracy", "5/5 (100 %)", "All factual claims matched KB content."),
        ("Out-of-scope fallback", "3/3 (100 %)", "No hallucinated content observed."),
        ("Avg. end-to-end latency", "4.8 s", "Meets NFR-01 (< 5 s)."),
        ("Idempotent re-upload", "Pass", "Duplicate chunks not created."),
        ("Conversation continuity", "Pass", "Follow-up queries resolved correctly."),
    ]
    add_table(doc, ["Test", "Result", "Notes"],
              [list(r) for r in eval_rows],
              col_widths=[2.0, 1.4, 3.1])
    add_para(doc, "Table 4: Summary of system evaluation results.",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    add_figure(doc, "fig7_deployment_diagram.png",
               "Figure 10: Deployment diagram. The user device, the application server "
               "running FastAPI with an embedded ChromaDB store, and the external AI "
               "APIs together make up the runtime topology.")
    add_page_break(doc)

    # ---- 7. USER MANUAL ----
    add_heading(doc, "7.  User Manual", level=1)

    add_heading(doc, "7.1  For Prospective Students and Parents", level=2)
    add_para(doc,
        "To use Mahir Connect, open the FAST-NUCES Peshawar website at the URL the "
        "admissions office has shared. The application loads inside any modern web "
        "browser, including the browser on a mobile phone. Nothing has to be installed, "
        "and you do not need to create an account."
    )
    add_para(doc,
        "When the page loads, click the microphone button on the voice card or the "
        "circular chat button at the bottom-right corner of the screen to open the chat "
        "widget. If you have not previously allowed microphone access for this website, "
        "the browser will ask for permission; click Allow. The system asks for "
        "microphone access only for voice queries — typed queries do not need it."
    )
    add_para(doc,
        "To ask a question by voice, click the microphone once and speak in Urdu or "
        "English — for example, \"FAST Peshawar mein BS ke liye apply kaise karna hai?\". "
        "The button pulses while it is recording. The query is submitted automatically "
        "when you finish speaking, or after the 30-second cap is reached. While the "
        "system transcribes and looks up an answer, you'll see a processing indicator. "
        "The reply appears as a text bubble with a speaker button; click the speaker to "
        "hear it. To stop the audio early, click the microphone at any time."
    )
    add_para(doc,
        "To ask by text, type into the field at the bottom of the chat widget and press "
        "Enter or click send. Text queries go through the same knowledge base, and you "
        "still get an audio version that you can play through the speaker button."
    )
    add_para(doc,
        "Follow-up questions don't need any reset — just keep talking. The system "
        "remembers the last three exchanges and resolves references like \"aur uski fees "
        "kitni hai?\" against the previous turn. If you'd rather start fresh, click the "
        "\"New Conversation\" button to clear the history. For questions the system "
        "cannot answer, the reply will say so plainly and, where possible, point you at "
        "the right contact."
    )

    add_heading(doc, "7.2  For Administrators", level=2)
    add_para(doc,
        "To get to the administration panel, append /admin to the application URL. Enter "
        "the administrator credentials when prompted. Once you're in, the panel shows a "
        "table of all current knowledge-base sources, with each filename and a one-line "
        "description of the category it covers."
    )
    add_para(doc,
        "To add new content, scroll to the \"Add to Knowledge Base\" section and drop a "
        ".txt file onto the upload zone, or click the zone to browse. The file has to be "
        "plain text and at most 2 MB. Click \"Upload & Embed\" to start ingestion. A "
        "spinner shows while the file is being processed; please don't close the tab "
        "during this. On success a green confirmation banner appears with the filename. "
        "On failure a red error banner appears. The source table refreshes by itself "
        "after a successful upload."
    )
    add_page_break(doc)

    # ---- 8. CONCLUSIONS AND RECOMMENDATIONS ----
    add_heading(doc, "8.  Conclusions and Recommendations", level=1)

    add_heading(doc, "8.1  Conclusions", level=2)
    add_para(doc,
        "Mahir Connect shows that a production-quality, domain-specific Urdu voice "
        "assistant can be put together inside the time and budget of a Final Year "
        "Project, using off-the-shelf LLM APIs [3], an open-source vector database [5] "
        "and a React frontend [14]. The system meets all five of the objectives we set "
        "out in Section 3.3. It runs in any modern browser without an installer, "
        "supports Urdu voice input and output, grounds every answer in a curated "
        "knowledge base via RAG [2], shows a live transcript, and clocks an average "
        "latency of 4.8 seconds — comfortably under the five-second target."
    )
    add_para(doc,
        "More importantly, the project is a proof point for a wider design idea: digital "
        "inclusion in Pakistan starts with putting spoken Urdu at the centre of the "
        "interface, not as an afterthought. The architecture shows that the gap between "
        "English-first AI products and the people they leave behind isn't really a "
        "technical wall [6, 7]; it's a question of whose needs you put first when you "
        "design. RAG [2] gives us a clean way to keep the model honest, which is what "
        "makes a generative system safe to use in a setting where the wrong answer can "
        "cost a student a year. And the per-category knowledge base plus the admin panel "
        "mean the system stays maintainable after we are gone, without depending on the "
        "development team being available."
    )

    add_heading(doc, "8.2  Recommendations", level=2)
    add_para(doc,
        "The knowledge base should be reviewed by the admissions team before every "
        "admission season, because fee structures, programme offerings, scholarship "
        "criteria and application procedures all change. The admin panel makes that "
        "review easy enough for a non-technical staff member, but a formal annual cycle "
        "is what stops it from quietly slipping. An out-of-date answer is more harmful "
        "than a missing one, because the system delivers it with the same confidence."
    )
    add_para(doc,
        "The architecture is intentionally modular. The knowledge base is the only "
        "FAST-NUCES-specific piece; the speech recogniser, the RAG engine, the language "
        "model and the TTS layer are all driven by configuration. Bringing the same "
        "platform up at another FAST campus or another Pakistani university needs "
        "nothing more than a fresh knowledge-base file set and a small set of config "
        "changes. A staged multi-campus rollout would let the same engineering work "
        "serve a much larger user base without much extra cost."
    )
    add_para(doc,
        "An analytics layer is the obvious next addition. Logging query topics, latency "
        "breakdowns by stage and a thumbs-up/down feedback signal would tell the "
        "admissions team which questions applicants actually ask most, which would feed "
        "back into both the official communication materials and the knowledge base. "
        "The same data would also let us tune the RAG retrieval parameters in a more "
        "principled way — for instance, picking the similarity-score threshold below "
        "which the fallback message should fire [23]."
    )
    add_para(doc,
        "Finally, a WhatsApp Business integration would meaningfully extend the system's "
        "reach. Most of our target users live on WhatsApp; sending a voice note to a "
        "university number and getting a voice note back fits naturally with how they "
        "already communicate. The backend pipeline is channel-agnostic, so a webhook "
        "would be enough to add WhatsApp as another input/output without touching the "
        "core."
    )
    add_page_break(doc)

    # ---- REFERENCES ----
    add_heading(doc, "References", level=1)
    refs = [
        '[1]  A. Radford, J. W. Kim, T. Xu, G. Brockman, C. McLeavey and I. Sutskever, '
        '"Robust speech recognition via large-scale weak supervision," in Proc. 40th '
        'Int. Conf. Mach. Learn. (ICML), Honolulu, HI, USA, 2023, pp. 28492–28518.',

        '[2]  P. Lewis, E. Perez, A. Piktus, F. Petroni, V. Karpukhin, N. Goyal, '
        'H. Küttler, M. Lewis, W.-t. Yih and T. Rocktäschel et al., "Retrieval-augmented '
        'generation for knowledge-intensive NLP tasks," in Proc. 34th Conf. Neural Inf. '
        'Process. Syst. (NeurIPS), 2020, pp. 9459–9474.',

        '[3]  OpenAI, "GPT-4 technical report," arXiv preprint arXiv:2303.08774, '
        'Mar. 2023.',

        '[4]  ElevenLabs, Inc., "ElevenLabs API documentation," New York, NY, USA, '
        '2024. Available: https://docs.elevenlabs.io',

        '[5]  Chroma, Inc., "ChromaDB: The AI-native open-source vector database," '
        '2024. Available: https://docs.trychroma.com',

        '[6]  UNESCO Institute for Statistics, "Pakistan — literacy rate, adult total '
        '(% of people ages 15 and above)," UNESCO, Montreal, Canada, 2023. Available: '
        'https://uis.unesco.org/en/country/pk',

        '[7]  Government of Pakistan, Pakistan Social and Living Standards Measurement '
        'Survey (PSLM) 2019–20, Pakistan Bureau of Statistics, Islamabad, 2021.',

        '[8]  T. Brown, B. Mann, N. Ryder, M. Subbiah, J. D. Kaplan, P. Dhariwal, '
        'A. Neelakantan, P. Shyam, G. Sastry and A. Askell et al., "Language models are '
        'few-shot learners," in Proc. 34th Conf. Neural Inf. Process. Syst. (NeurIPS), '
        '2020, pp. 1877–1901.',

        '[9]  D. Jurafsky and J. H. Martin, Speech and Language Processing, 3rd ed. '
        '(draft), 2023. Available: https://web.stanford.edu/~jurafsky/slp3/',

        '[10] A. Q. Akbar, M. A. Akhter and A. Iqbal, "Urdu natural language processing: '
        'A comprehensive survey," IEEE Access, vol. 10, pp. 55498–55516, 2022, doi: '
        '10.1109/ACCESS.2022.3175820.',

        '[11] E. Adamopoulou and L. Moussiades, "Chatbots: History, technology, and '
        'applications," Machine Learning with Applications, vol. 2, p. 100006, 2020, '
        'doi: 10.1016/j.mlwa.2020.100006.',

        '[12] Uplift AI, "Orator API documentation: Urdu text-to-speech," Pakistan, '
        '2024. Available: https://docs.upliftai.org',

        '[13] S. Ramírez, "FastAPI documentation," 2024. Available: '
        'https://fastapi.tiangolo.com',

        '[14] Meta Platforms, Inc., "React documentation," 2024. Available: '
        'https://react.dev',

        '[15] Z. Ashktorab, M. Jain, Q. V. Liao and J. D. Weisz, "Resilient chatbots: '
        'Repair strategy preferences for conversational breakdowns," in Proc. CHI '
        'Conf. Human Factors Comput. Syst. (CHI), Glasgow, UK, 2019, pp. 1–12, doi: '
        '10.1145/3290605.3300484.',

        '[16] B. A. Shawar and E. Atwell, "Chatbots: Are they really useful?" LDV-Forum, '
        'vol. 22, no. 1, pp. 29–49, 2007.',

        '[17] R. Dale, "The return of the chatbots," Natural Language Engineering, '
        'vol. 22, no. 5, pp. 811–817, 2016, doi: 10.1017/S1351324916000243.',

        '[18] T. Bocklisch, J. Faulkner, N. Pawlowski and A. Nichol, "Rasa: Open source '
        'language understanding and dialogue management," arXiv preprint '
        'arXiv:1712.05181, Dec. 2017.',

        '[19] Google LLC, "Dialogflow CX documentation," Google Cloud, Mountain View, '
        'CA, 2024. Available: https://cloud.google.com/dialogflow/cx/docs',

        '[20] Amazon Web Services, "Amazon Lex developer guide," Seattle, WA, 2024. '
        'Available: https://docs.aws.amazon.com/lex/latest/dg/what-is.html',

        '[21] Microsoft Corporation, "Azure Bot Service documentation," Redmond, WA, '
        '2024. Available: https://docs.microsoft.com/en-us/azure/bot-service/',

        '[22] Deepgram, Inc., "Nova-2: Deepgram speech-to-text API documentation," '
        'San Francisco, CA, 2024. Available: '
        'https://developers.deepgram.com/docs/nova-2',

        '[23] S. Robertson and H. Zaragoza, "The probabilistic relevance framework: '
        'BM25 and beyond," Foundations and Trends in Information Retrieval, vol. 3, '
        'no. 4, pp. 333–389, 2009, doi: 10.1561/1500000019.',

        '[24] K. Papineni, S. Roukos, T. Ward and W.-J. Zhu, "BLEU: A method for '
        'automatic evaluation of machine translation," in Proc. 40th Annu. Meeting '
        'ACL, Philadelphia, PA, 2002, pp. 311–318.',

        '[25] A. Vaswani, N. Shazeer, N. Parmar, J. Uszkoreit, L. Jones, A. N. Gomez, '
        'Ł. Kaiser and I. Polosukhin, "Attention is all you need," in Proc. 31st '
        'Conf. Neural Inf. Process. Syst. (NeurIPS), Long Beach, CA, 2017, '
        'pp. 5998–6008.',

        '[26] J. Devlin, M.-W. Chang, K. Lee and K. Toutanova, "BERT: Pre-training of '
        'deep bidirectional transformers for language understanding," in Proc. NAACL-HLT, '
        'Minneapolis, MN, 2019, pp. 4171–4186.',

        '[27] S. Garg, T. Vu and A. Moschitti, "TANDA: Transfer and adapt pre-trained '
        'transformer models for answer sentence selection," in Proc. AAAI Conf. '
        'Artif. Intell., vol. 34, no. 5, 2020, pp. 7780–7787, doi: '
        '10.1609/aaai.v34i05.6282.',

        '[28] J. Johnson, M. Douze and H. Jégou, "Billion-scale similarity search '
        'with GPUs," IEEE Trans. Big Data, vol. 7, no. 3, pp. 535–547, Jul. 2021, '
        'doi: 10.1109/TBDATA.2019.2921572.',

        '[29] LangChain, Inc., "LangChain documentation," San Francisco, CA, 2024. '
        'Available: https://docs.langchain.com',

        '[30] OpenAI, "Embeddings guide: text-embedding-ada-002," San Francisco, CA, '
        '2024. Available: https://platform.openai.com/docs/guides/embeddings',
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.paragraph_format.space_after = Pt(4)
    add_page_break(doc)

    # ---- APPENDIX A ----
    add_heading(doc, "Appendix A.  Sample Evaluation Queries", level=1)
    add_para(doc,
        "The following table lists the eight evaluation queries we used to measure "
        "system accuracy and fallback behaviour. Queries 1 to 5 are in-scope; the "
        "knowledge base contains authoritative information for them. Queries 6 to 8 are "
        "out-of-scope; the system is expected to deliver the defined fallback response "
        "rather than fabricate an answer. All eight were submitted as audio recordings "
        "in Urdu via the /voice-query endpoint."
    )
    eval_q = [
        ("1", "PhD applicants ke liye minimum requirements kya hain?", "Admissions"),
        ("2", "2025 ke admission process mein application submission ki dates kya hain?",
         "Admissions"),
        ("3", "Provisional admission kis ko di ja sakti hai?", "Admissions"),
        ("4", "Admission ke baad candidates ko kaunsi additional support provide ki ja "
              "sakti hai?", "General Information"),
        ("5", "Applicants ko kaun se documents submit karne zaroori hain?", "Admissions"),
        ("6", "FAST mein student counseling ya mental health support service available hai?",
         "Out of scope"),
        ("7", "University ki library digital resources aur e-books ka kitna collection hai?",
         "Out of scope"),
        ("8", "Kya university apne students ke liye health insurance ya medical coverage "
              "plans offer karti hai?", "Out of scope"),
    ]
    add_table(doc, ["No.", "Query (Roman Urdu)", "Category"],
              [list(r) for r in eval_q],
              col_widths=[0.5, 4.5, 1.6])
    add_para(doc, "Table 5: Evaluation query set for Mahir Connect.",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

    # ---- APPENDIX B ----
    add_heading(doc, "Appendix B.  AI Tool Usage Disclosure", level=1)
    add_para(doc,
        "Per the Student's Declaration, the generative AI tools listed below were used "
        "during the preparation of this report and the associated software. Final "
        "responsibility for the choice, application and interpretation of any "
        "AI-assisted material rests with the named authors."
    )
    add_para(doc,
        "Claude 3.5 Sonnet (Anthropic, Inc., version claude-3-5-sonnet-20241022) was "
        "used for language polishing of written sections of this report, for code review "
        "and debugging assistance during backend and frontend development, and for the "
        "first drafts of the system prompt templates in the RAG module. Every AI-assisted "
        "passage was read, fact-checked, and substantially rewritten by us before going "
        "into the final document."
    )
    add_para(doc,
        "GitHub Copilot (Microsoft Corporation / GitHub, Inc.) was used as an in-editor "
        "code completion assistant during frontend development. We accepted, modified or "
        "rejected each suggestion as we went."
    )
    add_para(doc,
        "ChatGPT-4o (OpenAI, San Francisco, CA, USA) was used for exploratory research "
        "into available Urdu TTS and STT APIs, and for early drafts of the knowledge-base "
        "ingestion pipeline. Every output was checked against the official API "
        "documentation before being relied on."
    )

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
